"""
Word Document Processor for Licia's Research Lab V2
Processes 30+ Word documents for editorial sprint
"""

import os
import json
import hashlib
from datetime import datetime
from typing import List, Dict, Any
import docx
import re
import asyncio

class WordDocumentProcessor:
    """Process Word documents for editorial sprint"""
    
    def __init__(self):
        self.processed_count = 0
        self.questions_extracted = []
        self.themes = {}
        self.insights = []
        
    def process_batch(self, doc_paths: List[str]) -> Dict[str, Any]:
        """
        Process 30+ Word documents for editorial sprint
        
        Args:
            doc_paths: List of paths to Word documents
            
        Returns:
            Dict containing questions, insights, themes, and raw content
        """
        results = {
            'questions': [],
            'insights': [],
            'themes': {},
            'raw_content': [],
            'metadata': {
                'total_documents': len(doc_paths),
                'processed_at': datetime.now().isoformat(),
                'processing_time': 0
            }
        }
        
        start_time = datetime.now()
        
        for path in doc_paths:
            try:
                # Process individual document
                doc_result = self.process_single_document(path)
                
                # Aggregate results
                results['questions'].extend(doc_result['questions'])
                results['insights'].extend(doc_result['insights'])
                results['raw_content'].append(doc_result['content'])
                
                # Merge themes
                for theme, count in doc_result['themes'].items():
                    results['themes'][theme] = results['themes'].get(theme, 0) + count
                
                self.processed_count += 1
                
            except Exception as e:
                print(f"Error processing {path}: {e}")
                results['metadata']['errors'] = results['metadata'].get('errors', [])
                results['metadata']['errors'].append({
                    'file': path,
                    'error': str(e)
                })
        
        # Calculate processing time
        results['metadata']['processing_time'] = (
            datetime.now() - start_time
        ).total_seconds()
        
        # Sort themes by frequency
        results['themes'] = dict(
            sorted(results['themes'].items(), key=lambda x: x[1], reverse=True)
        )
        
        # Cluster questions by theme
        results['question_clusters'] = self.cluster_questions(results['questions'])
        
        return results
    
    def process_single_document(self, doc_path: str) -> Dict[str, Any]:
        """Process a single Word document"""
        doc = docx.Document(doc_path)
        
        # Extract all text
        full_text = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)
        
        content = '\n'.join(full_text)
        
        # Extract questions
        questions = self.extract_questions(content)
        
        # Extract insights (lines with key insight markers)
        insights = self.extract_insights(content)
        
        # Extract themes
        themes = self.extract_themes(content)
        
        return {
            'source': os.path.basename(doc_path),
            'content': {
                'text': content,
                'hash': hashlib.sha256(content.encode()).hexdigest()[:8],
                'length': len(content),
                'source_file': doc_path
            },
            'questions': questions,
            'insights': insights,
            'themes': themes
        }
    
    def extract_questions(self, text: str) -> List[Dict[str, str]]:
        """Extract all questions from text"""
        questions = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if line.endswith('?'):
                # Get context (previous and next line if available)
                context_before = lines[i-1].strip() if i > 0 else ""
                context_after = lines[i+1].strip() if i < len(lines)-1 else ""
                
                questions.append({
                    'question': line,
                    'context_before': context_before,
                    'context_after': context_after,
                    'category': self.categorize_question(line)
                })
        
        return questions
    
    def extract_insights(self, text: str) -> List[Dict[str, str]]:
        """Extract key insights from text"""
        insights = []
        
        # Patterns that indicate insights
        insight_patterns = [
            r'(?i)key (?:insight|finding|discovery):?\s*(.+)',
            r'(?i)important:?\s*(.+)',
            r'(?i)note:?\s*(.+)',
            r'(?i)conclusion:?\s*(.+)',
            r'(?i)the (?:main|key|central) (?:point|idea|theme) is\s*(.+)',
            r'(?i)this (?:shows|demonstrates|reveals|suggests)\s*(.+)'
        ]
        
        for pattern in insight_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                insights.append({
                    'insight': match.strip(),
                    'type': pattern.split('(')[0].strip()
                })
        
        return insights
    
    def extract_themes(self, text: str) -> Dict[str, int]:
        """Extract and count themes from text"""
        themes = {}
        
        # Therapeutic themes to look for
        theme_keywords = {
            'trauma': ['trauma', 'traumatic', 'ptsd', 'stress disorder'],
            'attachment': ['attachment', 'bonding', 'secure base', 'safe haven'],
            'somatic': ['body', 'somatic', 'sensation', 'embodied', 'felt sense'],
            'emotional': ['emotion', 'feeling', 'affect', 'mood'],
            'safety': ['safety', 'safe', 'secure', 'trust'],
            'healing': ['heal', 'recovery', 'restoration', 'repair'],
            'nervous_system': ['nervous system', 'vagal', 'arousal', 'regulation'],
            'relationship': ['relationship', 'connection', 'alliance', 'rapport'],
            'integration': ['integration', 'integrate', 'synthesis', 'coherence'],
            'psychedelic': ['psychedelic', 'psilocybin', 'mdma', 'ketamine', 'integration'],
            'touch': ['touch', 'contact', 'proximity', 'tactile'],
            'senses': ['sense', 'sensory', 'perception', 'awareness']
        }
        
        lower_text = text.lower()
        
        for theme, keywords in theme_keywords.items():
            count = 0
            for keyword in keywords:
                count += lower_text.count(keyword)
            if count > 0:
                themes[theme] = count
        
        return themes
    
    def cluster_questions(self, questions: List[Dict]) -> Dict[str, List[Dict]]:
        """Cluster questions by theme for chapter organization"""
        clusters = {
            'foundational': [],
            'somatic': [],
            'relational': [],
            'integration': [],
            'practical': [],
            'research': [],
            'uncategorized': []
        }
        
        for q in questions:
            question_text = q['question'].lower()
            
            if any(word in question_text for word in ['what is', 'define', 'explain', 'foundation']):
                clusters['foundational'].append(q)
            elif any(word in question_text for word in ['body', 'somatic', 'sensation', 'feel']):
                clusters['somatic'].append(q)
            elif any(word in question_text for word in ['relationship', 'attachment', 'connect']):
                clusters['relational'].append(q)
            elif any(word in question_text for word in ['integrate', 'synthesis', 'combine']):
                clusters['integration'].append(q)
            elif any(word in question_text for word in ['how to', 'practice', 'apply', 'technique']):
                clusters['practical'].append(q)
            elif any(word in question_text for word in ['research', 'study', 'evidence', 'data']):
                clusters['research'].append(q)
            else:
                clusters['uncategorized'].append(q)
        
        # Remove empty clusters
        clusters = {k: v for k, v in clusters.items() if v}
        
        return clusters
    
    def categorize_question(self, question: str) -> str:
        """Categorize a single question"""
        q_lower = question.lower()
        
        if any(word in q_lower for word in ['what is', 'define', 'explain']):
            return 'definitional'
        elif any(word in q_lower for word in ['how', 'technique', 'method']):
            return 'procedural'
        elif any(word in q_lower for word in ['why', 'reason', 'cause']):
            return 'explanatory'
        elif any(word in q_lower for word in ['when', 'timing', 'sequence']):
            return 'temporal'
        elif any(word in q_lower for word in ['who', 'whom', 'population']):
            return 'demographic'
        else:
            return 'open-ended'
    
    async def process_batch_async(self, doc_paths: List[str]) -> Dict[str, Any]:
        """Process documents asynchronously for better performance"""
        tasks = []
        
        for path in doc_paths:
            task = asyncio.create_task(self.process_single_async(path))
            tasks.append(task)
        
        results = await asyncio.gather(*tasks)
        
        # Aggregate results
        return self.aggregate_results(results)
    
    async def process_single_async(self, doc_path: str) -> Dict:
        """Async wrapper for single document processing"""
        return await asyncio.to_thread(self.process_single_document, doc_path)
    
    def aggregate_results(self, results: List[Dict]) -> Dict:
        """Aggregate results from multiple documents"""
        aggregated = {
            'questions': [],
            'insights': [],
            'themes': {},
            'raw_content': [],
            'question_clusters': {}
        }
        
        for result in results:
            aggregated['questions'].extend(result['questions'])
            aggregated['insights'].extend(result['insights'])
            aggregated['raw_content'].append(result['content'])
            
            # Merge themes
            for theme, count in result['themes'].items():
                aggregated['themes'][theme] = aggregated['themes'].get(theme, 0) + count
        
        # Cluster all questions
        aggregated['question_clusters'] = self.cluster_questions(aggregated['questions'])
        
        return aggregated


def main():
    """Main entry point for processing Word documents"""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python word_processor.py <doc1.docx> [doc2.docx] ...")
        sys.exit(1)
    
    doc_paths = sys.argv[1:]
    processor = WordDocumentProcessor()
    
    print(f"Processing {len(doc_paths)} Word documents...")
    results = processor.process_batch(doc_paths)
    
    # Save results
    output_file = 'outputs/word_processing_results.json'
    os.makedirs('outputs', exist_ok=True)
    
    with open(output_file, 'w') as f:
        json.dump(results, f, indent=2)
    
    print(f"\nProcessing complete!")
    print(f"- Documents processed: {processor.processed_count}")
    print(f"- Questions extracted: {len(results['questions'])}")
    print(f"- Insights found: {len(results['insights'])}")
    print(f"- Themes identified: {len(results['themes'])}")
    print(f"\nResults saved to: {output_file}")
    
    # Show question clusters
    print("\nQuestion Clusters for Chapters:")
    for cluster_name, questions in results['question_clusters'].items():
        print(f"  {cluster_name}: {len(questions)} questions")


if __name__ == "__main__":
    main()
